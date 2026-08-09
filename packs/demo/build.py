"""Generate the demo pack — five complete, valid supplier quotations.

    python packs/demo/build.py

Every supplier here is clean. Nothing is missing, nothing contradicts, no document is
unreadable. The ranking is decided entirely by commercial values, which is the point:
this pack answers "which supplier should we go with", not "which supplier filled the
form in wrong". Defect handling has its own fixtures in `packs/synthetic`.

    1  Zhongshan Polymer Works   China       FOB Shenzhen    3.85/pc   tooling 14,500
    2  Pune Precision Polymers   India       DDP Karachi     4.55/pc   tooling  9,500
    3  PT Batam Injection        Indonesia   FOB Batam       3.42/pc   tooling 38,000
    4  Bac Ninh Moulding JSC     Vietnam     CIF Karachi     4.10/pc   tooling 16,000
    5  Konya Kalip Sanayi        Turkiye     EXW Konya       3.75/pc   tooling 12,000

The economics are built so the answer cannot be read off the quotations:

  - **Pune wins at 20,000 with the highest unit price on the table.** DDP puts main
    carriage, insurance and import duty on the seller, so 4.55 delivered beats 3.85
    at the factory gate. A buyer comparing unit prices picks the wrong supplier —
    which is the entire thesis of a landed-cost tool.
  - **Batam has the cheapest unit price and finishes last**, because 38,000 of tooling
    spread over 20,000 pieces is 1.90 a unit. Re-run at 100,000 and Batam takes first
    place. The right answer depends on the volume, which is why the app asks for it.
  - **Konya is EXW**, the term under which the buyer carries every leg. Its low unit
    price is not the bargain it appears to be.

Formats span what the ingest layer must handle: PDF, DOCX, XLSX, CSV.

Deterministic: no randomness, every embedded timestamp pinned, so a rebuild is
byte-identical and a diff here is always a change someone made on purpose.

Authored by us. Not organizer data. Never used for any reported metric.
"""

from __future__ import annotations

import csv
import re
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document as Docx
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

HERE = Path(__file__).parent
PAGE_WIDTH, PAGE_HEIGHT = A4
PINNED_TIME = datetime(2026, 8, 12, 9, 0, 0)

CORE_PROPERTIES = "docProps/core.xml"
MODIFIED_TAG = re.compile(rb"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)")

PART = "CTR-H220"
PART_DESC = "Controller housing, ABS+PC, textured black, ultrasonic-weld ready"
BUYER = "Meridian Controls Ltd."
ENQUIRY = "MC-RFQ-2026-041"
VOLUME = 20_000


# --------------------------------------------------------------------------- #
# Determinism
# --------------------------------------------------------------------------- #

def pin_office_timestamps(path: Path) -> None:
    """Strip the two clock readings an Office file records on save.

    DOCX and XLSX are ZIP archives whose entries carry the wall-clock time they were
    written, and openpyxl rewrites `dcterms:modified` regardless of what was set on
    the workbook. Both are normalised so rebuilds stay byte-identical.
    """
    stamp = PINNED_TIME.timetuple()[:6]
    iso = PINNED_TIME.strftime("%Y-%m-%dT%H:%M:%SZ").encode()

    with zipfile.ZipFile(path) as source:
        entries = [(info, source.read(info.filename)) for info in source.infolist()]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as target:
        for info, data in entries:
            if info.filename == CORE_PROPERTIES:
                data = MODIFIED_TAG.sub(rb"\g<1>" + iso + rb"\g<2>", data)
            pinned = zipfile.ZipInfo(info.filename, date_time=stamp)
            pinned.compress_type = info.compress_type
            pinned.external_attr = info.external_attr
            target.writestr(pinned, data)


def save_docx(document: Docx, path: Path) -> None:
    document.core_properties.created = PINNED_TIME
    document.core_properties.modified = PINNED_TIME
    document.save(path)
    pin_office_timestamps(path)


# --------------------------------------------------------------------------- #
# Terms every quotation states, in one place
# --------------------------------------------------------------------------- #

def terms(
    unit_price: str,
    tooling: str,
    moq: str,
    lead_time: str,
    weight: str,
    delivery: str,
    payment: str,
    origin: str,
) -> list[tuple[str, str]]:
    """The commercial block, in a fixed order and with fixed labels.

    Same labels on every quotation on purpose. Real suppliers vary their wording, and
    the synthetic pack exercises that; here the variable under test is the economics,
    so everything else is held constant.
    """
    return [
        ("Unit price", unit_price),
        ("Price basis", "Per piece, ex-tooling"),
        ("Tooling charge", tooling),
        ("Minimum order quantity", moq),
        ("Production lead time", lead_time),
        ("Unit weight", weight),
        ("Delivery terms", delivery),
        ("Payment terms", payment),
        ("Country of origin", origin),
        ("HS code", "8538.90"),
    ]


# --------------------------------------------------------------------------- #
# PDF quotation layout
# --------------------------------------------------------------------------- #

LEFT = 20 * mm
RIGHT = PAGE_WIDTH - 20 * mm


class Sheet:
    """A one-page quotation on letterhead, with a real text layer."""

    def __init__(self, path: Path, title: str) -> None:
        self.page = canvas.Canvas(str(path), pagesize=A4, invariant=1)
        self.page.setTitle(title)
        self.y = PAGE_HEIGHT - 20 * mm

    def letterhead(self, name: str, address: list[str], contact: str) -> None:
        self.page.setFont("Helvetica-Bold", 15)
        self.page.drawString(LEFT, self.y, name)
        self.y -= 5.5 * mm
        self.page.setFont("Helvetica", 8.5)
        for line in address:
            self.page.drawString(LEFT, self.y, line)
            self.y -= 4 * mm
        self.page.drawString(LEFT, self.y, contact)
        self.y -= 6 * mm
        self.rule()

    def rule(self, gap: float = 5.0) -> None:
        self.page.setLineWidth(0.6)
        self.page.line(LEFT, self.y, RIGHT, self.y)
        self.y -= gap * mm

    def heading(self, text: str) -> None:
        self.page.setFont("Helvetica-Bold", 12)
        self.page.drawString(LEFT, self.y, text)
        self.y -= 6 * mm

    def meta(self, pairs: list[tuple[str, str]]) -> None:
        for index in range(0, len(pairs), 2):
            for column, (label, value) in enumerate(pairs[index:index + 2]):
                x = LEFT + column * 85 * mm
                self.page.setFont("Helvetica-Bold", 9)
                self.page.drawString(x, self.y, f"{label}:")
                self.page.setFont("Helvetica", 9)
                self.page.drawString(x + 32 * mm, self.y, value)
            self.y -= 5 * mm
        self.y -= 2 * mm

    def table(self, rows: list[list[str]]) -> None:
        columns = [("Item", 14, "l"), ("Part / Description", 76, "l"),
                   ("Qty", 22, "r"), ("Unit Price", 30, "r"), ("Amount", 28, "r")]
        self.page.setFont("Helvetica-Bold", 8.5)
        x = LEFT
        for header, width, align in columns:
            draw = self.page.drawRightString if align == "r" else self.page.drawString
            draw(x + (width * mm if align == "r" else 0), self.y, header)
            x += width * mm
        self.y -= 2 * mm
        self.rule(gap=4.0)

        self.page.setFont("Helvetica", 9)
        for row in rows:
            x = LEFT
            for value, (_, width, align) in zip(row, columns, strict=True):
                draw = self.page.drawRightString if align == "r" else self.page.drawString
                draw(x + (width * mm if align == "r" else 0), self.y, value)
                x += width * mm
            self.y -= 5 * mm
        self.y -= 1 * mm
        self.rule(gap=6.0)

    def block(self, title: str, pairs: list[tuple[str, str]]) -> None:
        self.page.setFont("Helvetica-Bold", 10)
        self.page.drawString(LEFT, self.y, title)
        self.y -= 5.5 * mm
        for label, value in pairs:
            self.page.setFont("Helvetica-Bold", 9)
            self.page.drawString(LEFT, self.y, f"{label}:")
            self.page.setFont("Helvetica", 9)
            self.page.drawString(LEFT + 45 * mm, self.y, value)
            self.y -= 5 * mm
        self.y -= 2 * mm

    def paragraph(self, title: str, lines: list[str]) -> None:
        self.page.setFont("Helvetica-Bold", 10)
        self.page.drawString(LEFT, self.y, title)
        self.y -= 5.5 * mm
        self.page.setFont("Helvetica", 8.5)
        for line in lines:
            self.page.drawString(LEFT, self.y, line)
            self.y -= 4.5 * mm
        self.y -= 3 * mm

    def signature(self, name: str, role: str) -> None:
        self.y -= 4 * mm
        self.page.setFont("Helvetica", 8.5)
        self.page.drawString(LEFT, self.y, "For and on behalf of the seller,")
        self.y -= 10 * mm
        self.page.setFont("Helvetica-Bold", 9)
        self.page.drawString(LEFT, self.y, name)
        self.y -= 4.5 * mm
        self.page.setFont("Helvetica", 8.5)
        self.page.drawString(LEFT, self.y, role)

    def save(self) -> None:
        self.page.showPage()
        self.page.save()


def pdf_quotation(
    filename: str,
    company: str,
    address: list[str],
    contact: str,
    ref: str,
    date: str,
    validity: str,
    line_items: list[list[str]],
    commercial: list[tuple[str, str]],
    notes: list[str],
    bank: list[str],
    signer: tuple[str, str],
    heading: str = "QUOTATION",
) -> None:
    sheet = Sheet(HERE / filename, ref)
    sheet.letterhead(company, address, contact)
    sheet.heading(heading)
    sheet.meta([
        ("Quotation No", ref), ("Date", date),
        ("Buyer", BUYER), ("Validity", validity),
        ("Enquiry ref", ENQUIRY), ("Currency", "USD"),
    ])
    sheet.table(line_items)
    sheet.block("COMMERCIAL TERMS", commercial)
    sheet.paragraph("NOTES", notes)
    sheet.paragraph("BANK DETAILS", bank)
    sheet.signature(*signer)
    sheet.save()


# --------------------------------------------------------------------------- #
# 1 — Zhongshan Polymer Works (China) — FOB
# --------------------------------------------------------------------------- #

def zhongshan() -> None:
    pdf_quotation(
        "Zhongshan Polymer Works - Quotation ZPW-2026-0812.pdf",
        "ZHONGSHAN POLYMER WORKS CO., LTD.",
        ["Building 7, Torch Development Zone, Zhongshan, Guangdong 528437, China",
         "Injection moulding | Tooling | Ultrasonic assembly | ISO 9001:2015"],
        "T +86 760 8832 4417   E sales@zspolymer.com.cn   Contact: Ms. Li Wenjing",
        "ZPW-2026-0812", "12 August 2026", "45 days from date of issue",
        [["1", f"{PART} - {PART_DESC}", "20,000 pcs", "USD 3.85", "USD 77,000.00"],
         ["2", "Injection tool, 2-cavity, P20 steel (one-off)", "1 set",
          "USD 14,500.00", "USD 14,500.00"]],
        terms("USD 3.85 per piece", "USD 14,500.00 one-off, non-refundable",
              "5,000 pieces", "35 days from tooling approval", "0.185 kg per piece",
              "FOB Shenzhen (Incoterms 2020)",
              "30% advance with PO, 70% against B/L copy", "China"),
        ["1. Tooling remains the property of the buyer on full settlement of the tooling charge.",
         "2. Material: ABS+PC Bayblend T65XF or approved equivalent, UL94 V-0.",
         "3. First article inspection report supplied with the initial shipment at no charge.",
         "4. Capacity reserved for this part: 120,000 pieces per month.",
         "5. Prices firm for the validity period; thereafter subject to resin index review."],
        ["Beneficiary: Zhongshan Polymer Works Co., Ltd.",
         "Bank: Bank of China, Zhongshan Branch    SWIFT: BKCHCNBJ440"],
        ("Li Wenjing", "Export Sales Manager"),
    )


# --------------------------------------------------------------------------- #
# 2 — Pune Precision Polymers (India) — DDP, and the winner at 20,000
# --------------------------------------------------------------------------- #

def pune() -> None:
    """A quotation as a Word document, which is how a good many of them arrive.

    Dearest per piece and cheapest once landed. DDP means the seller clears the goods
    into the destination country, so freight, insurance and duty are already inside
    that 4.55 — the three lines that are added on top of everybody else's price.
    """
    document = Docx()
    document.add_heading("PUNE PRECISION POLYMERS PVT. LTD.", level=1)
    document.add_paragraph(
        "Plot 44, Chakan Industrial Area Phase II, Pune, Maharashtra 410501, India"
    )
    document.add_paragraph(
        "T +91 20 6741 2200 | E exports@puneprecision.co.in | Contact: Mr. Rohit Deshmukh"
    )
    document.add_paragraph("Injection moulding | Tool room | ISO 9001:2015, IATF 16949")

    document.add_heading("QUOTATION", level=2)
    for label, value in [
        ("Quotation No", "PPP/EXP/2026/0447"),
        ("Date", "13 August 2026"),
        ("Buyer", BUYER),
        ("Enquiry reference", ENQUIRY),
        ("Validity", "60 days from date of issue"),
        ("Currency", "USD"),
    ]:
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Line items", level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, header in zip(
        table.rows[0].cells,
        ["Item", "Part / Description", "Quantity", "Unit price", "Amount"],
        strict=True,
    ):
        cell.text = header
    for values in [
        ["1", f"{PART} - {PART_DESC}", "20,000 pcs", "USD 4.55", "USD 91,000.00"],
        ["2", "Injection tool, 2-cavity, one-off", "1 set", "USD 9,500.00", "USD 9,500.00"],
    ]:
        for cell, value in zip(table.add_row().cells, values, strict=True):
            cell.text = value

    document.add_heading("Commercial terms", level=2)
    for label, value in terms(
        "USD 4.55 per piece", "USD 9,500.00 one-off", "8,000 pieces",
        "42 days from tooling approval", "0.188 kg per piece",
        "DDP Karachi (Incoterms 2020)",
        "T/T 25% advance, 75% net 60 days from invoice", "India",
    ):
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Notes", level=2)
    for line in [
        "1. DDP price is inclusive of ocean freight, marine insurance, import duty and "
        "customs clearance at Karachi. No further landed charges are payable by the buyer.",
        "2. Tooling is transferred to the buyer on settlement of the tooling charge.",
        "3. Material: ABS+PC, UL94 V-0, RoHS and REACH compliant.",
        "4. Monthly capacity available for this part: 150,000 pieces.",
        "5. PPAP Level 3 documentation supplied prior to series release.",
    ]:
        document.add_paragraph(line)

    document.add_paragraph("Beneficiary: Pune Precision Polymers Pvt. Ltd.")
    document.add_paragraph("Bank: HDFC Bank, Chakan Branch    SWIFT: HDFCINBB")
    document.add_paragraph("For and on behalf of the seller,")
    document.add_paragraph("Rohit Deshmukh, Head of Exports")

    save_docx(document, HERE / "Pune Precision Polymers - Quotation PPP-EXP-2026-0447.docx")


# --------------------------------------------------------------------------- #
# 3 — PT Batam Injection Molding (Indonesia) — cheapest per piece, worst landed
# --------------------------------------------------------------------------- #

HEADER_FILL = PatternFill("solid", fgColor="DDE5F0")
THIN = Side(style="thin", color="99A3B0")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def batam() -> None:
    """The quotation as a worksheet.

    3.42 a piece is the lowest number in the pack and the reason this supplier loses:
    38,000 of tooling over 20,000 pieces adds 1.90 a unit. At 100,000 the same tooling
    adds 0.38 and Batam wins outright. Nothing about that is visible on the quotation.
    """
    book = Workbook()
    sheet = book.active
    sheet.title = "Quotation"

    for label, value in [
        ("PT BATAM INJECTION MOLDING", None),
        ("Kawasan Industri Batamindo, Muka Kuning, Batam 29433, Indonesia", None),
        ("T +62 778 611 340 | E sales@bataminjection.co.id", None),
        (None, None),
        ("QUOTATION", None),
        ("Quotation No", "BIM/Q/2026/0812"),
        ("Date", "12 August 2026"),
        ("Buyer", BUYER),
        ("Enquiry reference", ENQUIRY),
        ("Validity", "30 days"),
        ("Currency", "USD"),
        (None, None),
        ("LINE ITEMS", None),
    ]:
        sheet.append([label, value])

    sheet.append(["Item", "Description", "Quantity", "Unit price (USD)", "Amount (USD)"])
    header_row = sheet.max_row
    sheet.append([1, f"{PART} - {PART_DESC}", VOLUME, 3.42, 68400.00])
    sheet.append([2, "Injection tool, 4-cavity hot runner, one-off", 1, 38000.00, 38000.00])

    sheet.append([None, None])
    sheet.append(["COMMERCIAL TERMS", None])
    for label, value in terms(
        "USD 3.42 per piece", "USD 38,000.00 one-off", "20,000 pieces",
        "45 days from tooling approval", "0.183 kg per piece",
        "FOB Batam (Incoterms 2020)",
        "T/T 35% advance, 65% against shipping documents", "Indonesia",
    ):
        sheet.append([label, value])

    sheet.append([None, None])
    sheet.append(["Note", "4-cavity hot runner tool. Higher tooling cost, lowest piece "
                          "price in class. Best value above 60,000 pieces per year."])
    sheet.append(["Capacity", "200,000 pieces per month"])
    sheet.append(["Certifications", "ISO 9001:2015, ISO 14001:2015"])
    sheet.append(["Prepared by", "Andi Kurniawan, Sales Manager"])

    for cell in sheet[header_row]:
        if cell.value is not None:
            cell.font = Font(bold=True)
            cell.fill = HEADER_FILL
            cell.border = BOX
    for row in (1, 5):
        sheet.cell(row=row, column=1).font = Font(bold=True, size=13 if row == 1 else 11)
    for column, width in zip("ABCDE", (26, 58, 12, 17, 16), strict=True):
        sheet.column_dimensions[column].width = width
    for column in "CDE":
        for cell in sheet[column]:
            cell.alignment = Alignment(horizontal="right")

    path = HERE / "PT Batam Injection - Quotation BIM-Q-2026-0812.xlsx"
    book.properties.created = PINNED_TIME
    book.properties.modified = PINNED_TIME
    book.save(path)
    pin_office_timestamps(path)


# --------------------------------------------------------------------------- #
# 4 — Bac Ninh Moulding JSC (Vietnam) — CIF
# --------------------------------------------------------------------------- #

def bac_ninh() -> None:
    pdf_quotation(
        "Bac Ninh Moulding - Quotation BNM-Q-2291.pdf",
        "BAC NINH MOULDING JSC",
        ["Lot C4, Que Vo Industrial Park, Bac Ninh Province, Vietnam",
         "Injection moulding | Cleanroom assembly | ISO 9001:2015, ISO 14001:2015"],
        "T +84 222 3634 880   E export@bacninhmoulding.vn   Contact: Ms. Nguyen Thi Hoa",
        "BNM-Q-2291", "14 August 2026", "30 days",
        [["1", f"{PART} - {PART_DESC}", "20,000 pcs", "USD 4.10", "USD 82,000.00"],
         ["2", "Injection tool, 2+2 cavity, one-off", "1 set",
          "USD 16,000.00", "USD 16,000.00"]],
        terms("USD 4.10 per piece", "USD 16,000.00 one-off", "10,000 pieces",
              "40 days from tooling approval", "0.190 kg per piece",
              "CIF Karachi (Incoterms 2020)",
              "T/T 40% advance, 60% against shipping documents", "Vietnam"),
        ["1. CIF price includes ocean freight and marine insurance to Karachi.",
         "   Import duty and customs clearance remain for the buyer's account.",
         "2. Material: ABS+PC, UL94 V-0.",
         "3. Cleanroom assembly cell available, ISO Class 8, commissioned 2024.",
         "4. Monthly capacity: 90,000 pieces for this product family."],
        ["Beneficiary: Bac Ninh Moulding JSC",
         "Bank: Vietcombank, Bac Ninh Branch    SWIFT: BFTVVNVX"],
        ("Nguyen Thi Hoa", "Export Sales Manager"),
    )


def bac_ninh_profile() -> None:
    """A capability document that agrees with the quotation.

    Present as supporting evidence, not as a contradiction — the point of this pack is
    that the decision turns on values, so the profile corroborates rather than disputes.
    """
    document = Docx()
    document.add_heading("Bac Ninh Moulding JSC - Company Profile 2026", level=1)
    document.add_paragraph(
        "Bac Ninh Moulding JSC is a tier-two injection moulding supplier serving "
        "consumer electronics and industrial controls customers across Europe and "
        "North America. Established 2011. 240 employees across two sites in Bac Ninh "
        "Province, Vietnam."
    )

    for heading, lines in [
        ("Manufacturing capability", [
            "18 injection moulding machines, 50 to 650 tonne clamping force.",
            "In-house tool room: 2+2 and 4-cavity tools, P20 and H13 steel.",
            "Secondary operations: pad printing, ultrasonic welding, heat staking.",
            "Cleanroom assembly cell, ISO Class 8, commissioned 2024.",
        ]),
        ("Commercial terms", [
            "Minimum order quantity: 10,000 pieces per part number.",
            "Typical production lead time: 40 days from tool sign-off.",
            "Monthly capacity: 90,000 pieces for this product family.",
            "Standard payment terms: T/T 40% advance, balance against documents.",
        ]),
        ("Certifications", [
            "ISO 9001:2015 - certificate VN-QMS-11482, valid to March 2028.",
            "ISO 14001:2015 - certificate VN-EMS-3391, valid to March 2028.",
            "IATF 16949 - in progress, audit scheduled Q4 2026.",
            "RoHS and REACH declarations available on request.",
        ]),
        ("Quality performance, rolling 12 months", [
            "On-time delivery: 94.2%.",
            "Parts per million defective: 830.",
            "Customer complaints closed within 30 days: 100%.",
        ]),
    ]:
        document.add_heading(heading, level=2)
        for line in lines:
            document.add_paragraph(line, style="List Bullet")

    document.add_paragraph(
        "Contact: Nguyen Thi Hoa, Export Sales - export@bacninhmoulding.vn"
    )
    save_docx(document, HERE / "Bac Ninh Moulding - Company Profile 2026.docx")


# --------------------------------------------------------------------------- #
# 5 — Konya Kalip Sanayi (Turkiye) — EXW
# --------------------------------------------------------------------------- #

def konya() -> None:
    """EXW: the buyer carries every leg from the factory gate onward.

    A low unit price under EXW is not the bargain it looks like, because everything
    Pune has already absorbed is still to be added here.
    """
    pdf_quotation(
        "Konya Kalip - Teklif KKS-2026-118.pdf",
        "KONYA KALIP SANAYI A.S.",
        ["Organize Sanayi Bolgesi, 8. Sokak No 14, Selcuklu, Konya 42300, Turkiye",
         "Plastic injection moulding and tool manufacture | ISO 9001:2015"],
        "T +90 332 239 1180   E ihracat@konyakalip.com.tr   Contact: Mr. Emre Dogan",
        "KKS-2026-118", "13 August 2026", "30 days",
        [["1", f"{PART} - {PART_DESC}", "20,000 pcs", "USD 3.75", "USD 75,000.00"],
         ["2", "Mould tool, 2-cavity, one-off charge", "1 set",
          "USD 12,000.00", "USD 12,000.00"]],
        terms("USD 3.75 per piece", "USD 12,000.00 one-off", "3,000 pieces",
              "28 days from tooling approval", "0.182 kg per piece",
              "EXW Konya (Incoterms 2020)",
              "50% with order, 50% prior to despatch", "Turkiye"),
        ["1. EXW: collection, export clearance, carriage, insurance and import duty",
         "   are for the buyer's account.",
         "2. Packing: 200 pieces per carton, palletised, 40 cartons per pallet.",
         "3. Material: ABS+PC, UL94 V-0.",
         "4. Monthly capacity: 70,000 pieces.",
         "5. Tooling lead time is additional to the production lead time stated above."],
        ["Beneficiary: Konya Kalip Sanayi A.S.",
         "Bank: Turkiye Is Bankasi, Konya OSB    SWIFT: ISBKTRIS"],
        ("Emre Dogan", "Foreign Trade Manager"),
        heading="QUOTATION / TEKLIF",
    )


# --------------------------------------------------------------------------- #
# Shared project material
# --------------------------------------------------------------------------- #

def product_brief() -> None:
    document = Docx()
    document.add_heading(f"Product Brief - {PART} Controller Housing", level=1)
    document.add_paragraph(
        f"Issued by {BUYER} for sourcing enquiry {ENQUIRY}. This brief accompanies the "
        "request for quotation and states the requirements every offer is measured "
        "against."
    )

    for heading, lines in [
        ("Requirement", [
            f"Part number: {PART}.",
            f"Description: {PART_DESC}.",
            f"Target annual volume: {VOLUME:,} pieces, released in two equal shipments.",
            "Destination: Karachi, Pakistan (buyer's contract assembler).",
            "Reporting currency for comparison: USD.",
        ]),
        ("Mandatory requirements", [
            "Material: ABS+PC blend, UL94 V-0 rated.",
            "Supplier must hold a current ISO 9001 certificate.",
            "Tooling must become the property of Meridian Controls on settlement.",
            "First article inspection report required before series release.",
        ]),
        ("What each quotation must state", [
            "Unit price and the basis it is quoted on.",
            "Currency.",
            "Delivery terms as an Incoterms 2020 rule, with the named place.",
            "One-off tooling charge, if any.",
            "Minimum order quantity and production lead time.",
            "Unit weight, for freight assessment.",
        ]),
        ("How offers will be compared", [
            "On landed cost per piece at the target volume, not on unit price.",
            "Delivery terms materially change who pays for carriage, insurance and duty,",
            "  so two identical unit prices can differ by thousands once landed.",
            "Tooling is amortised across the order, so the ranking is volume-dependent",
            "  and will be re-run if the volume changes.",
        ]),
    ]:
        document.add_heading(heading, level=2)
        for line in lines:
            document.add_paragraph(line, style="List Bullet")

    save_docx(document, HERE / f"Product Brief - {PART}.docx")


def bill_of_materials() -> None:
    rows = [
        ["Level", "Item", "Description", "Material", "Qty per unit", "Unit"],
        ["1", f"{PART}-A", "Housing, upper shell", "ABS+PC UL94 V-0", "1", "pc"],
        ["1", f"{PART}-B", "Housing, lower shell", "ABS+PC UL94 V-0", "1", "pc"],
        ["2", f"{PART}-G", "Gasket, perimeter seal", "Silicone 60 Shore A", "1", "pc"],
        ["2", f"{PART}-L", "Light pipe, status LED", "PMMA clear", "2", "pc"],
        ["2", "FST-M3-08", "Screw, M3 x 8 pan head", "Stainless A2", "4", "pc"],
        ["3", f"LBL-{PART[-4:]}", "Rating label, printed", "Polyester", "1", "pc"],
    ]
    with (HERE / f"Bill of Materials - {PART}.csv").open(
        "w", newline="", encoding="utf-8"
    ) as handle:
        csv.writer(handle).writerows(rows)


# --------------------------------------------------------------------------- #

def main() -> None:
    HERE.mkdir(parents=True, exist_ok=True)
    for stale in HERE.glob("*.xlsx"):
        stale.unlink()
    for stale in HERE.glob("*.pdf"):
        stale.unlink()

    zhongshan()
    pune()
    batam()
    bac_ninh()
    bac_ninh_profile()
    konya()
    product_brief()
    bill_of_materials()

    for path in sorted(HERE.iterdir()):
        if path.suffix.lower() in {".pdf", ".xlsx", ".docx", ".csv"}:
            print(f"  {path.name:<58} {path.stat().st_size:>7,} bytes")


if __name__ == "__main__":
    main()
