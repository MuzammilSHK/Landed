"""Documents in, positioned text out.

Deliberately dumb: this layer performs no interpretation whatsoever. It turns bytes
into text that still knows which page, sheet, and cell it came from, because every
citation in the finished product traces back to an anchor produced here. Losing
position information at this stage silently breaks provenance everywhere downstream.

A page with no meaningful text layer is not a failure — it is a scan. Those pages are
rendered to images and handed to the vision path, with `ReadMethod.VISION` recorded
so the value they yield is marked for verification rather than trusted equally.

Unreadable files are reported, never skipped. A quotation that silently failed to
parse is a supplier missing from the comparison, which is worse than an error.
"""

from __future__ import annotations

import csv
import hashlib
import mimetypes
from io import BytesIO
from pathlib import Path

import pdfplumber
import pypdfium2 as pdfium
from docx import Document as Docx
from openpyxl import load_workbook
from PIL import Image
from pydantic import BaseModel, Field

from .schema import ReadMethod

# Below this many characters, a PDF page is a scan with incidental artefacts rather
# than a page we can read. Page numbers and stray marks routinely yield a handful.
SCANNED_TEXT_THRESHOLD = 24

# Rendering scale for scanned pages. 2.0 is roughly 144 DPI — enough for a model to
# read quotation tables without producing payloads that cost more than they inform.
RENDER_SCALE = 2.0
MAX_IMAGE_EDGE = 2000

PDF_SUFFIXES = {".pdf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp", ".bmp"}
DOCX_SUFFIXES = {".docx"}
EXCEL_SUFFIXES = {".xlsx", ".xlsm"}
DELIMITED_SUFFIXES = {".csv", ".tsv"}
TEXT_SUFFIXES = {".txt", ".md"}


class Chunk(BaseModel):
    """A positioned piece of text from a source document."""

    text: str
    file: str
    page: int | None = None
    sheet: str | None = None
    read_method: ReadMethod = ReadMethod.TEXT_LAYER


class PageImage(BaseModel):
    """A page with no text layer, rendered for the vision path."""

    data: bytes
    media_type: str
    file: str
    page: int


class IngestedDocument(BaseModel):
    """Everything the extraction layer needs about one file.

    `error` is populated instead of raising: one unreadable quotation should not
    abort a pack, but it must not vanish either.
    """

    filename: str
    sha256: str
    byte_size: int
    content_type: str | None = None
    chunks: list[Chunk] = Field(default_factory=list)
    images: list[PageImage] = Field(default_factory=list)
    error: str | None = None

    @property
    def is_readable(self) -> bool:
        return self.error is None and bool(self.chunks or self.images)

    @property
    def needs_vision(self) -> bool:
        return bool(self.images)


def ingest_file(path: Path) -> IngestedDocument:
    """Read one document. Never raises for content reasons — see `error`."""
    raw = path.read_bytes()
    document = IngestedDocument(
        filename=path.name,
        sha256=hashlib.sha256(raw).hexdigest(),
        byte_size=len(raw),
        content_type=mimetypes.guess_type(path.name)[0],
    )

    suffix = path.suffix.lower()
    loader = _loader_for(suffix)
    if loader is None:
        return document.model_copy(
            update={"error": f"unsupported file type '{suffix or path.name}'"}
        )
    try:
        chunks, images = loader(path)
    except Exception as exc:
        # Broad on purpose: parser libraries raise their own hierarchies, and any
        # failure here is a data gap the user must see rather than a crash.
        return document.model_copy(
            update={"error": f"{type(exc).__name__}: {exc}"}
        )
    if not chunks and not images:
        return document.model_copy(update={"error": "no readable content found"})
    return document.model_copy(update={"chunks": chunks, "images": images})


def ingest_pack(directory: Path) -> list[IngestedDocument]:
    """Read every file in a pack directory, in a stable order.

    Sorted so a re-run produces the same sequence — extraction costs money and
    reproducibility matters more than directory order.
    """
    return [
        ingest_file(path)
        for path in sorted(directory.iterdir())
        if path.is_file() and path.suffix.lower() not in {".py", ".md", ".json"}
    ]


def _loader_for(suffix: str):
    if suffix in PDF_SUFFIXES:
        return _load_pdf
    if suffix in IMAGE_SUFFIXES:
        return _load_image
    if suffix in DOCX_SUFFIXES:
        return _load_docx
    if suffix in EXCEL_SUFFIXES:
        return _load_excel
    if suffix in DELIMITED_SUFFIXES:
        return _load_delimited
    if suffix in TEXT_SUFFIXES:
        return _load_text
    return None


def _load_pdf(path: Path) -> tuple[list[Chunk], list[PageImage]]:
    """One chunk per readable page; scanned pages are rendered instead.

    Mixed documents are common — a typed quotation with a photographed certificate
    stapled on — so the decision is made per page, not per file.
    """
    chunks: list[Chunk] = []
    scanned: list[int] = []
    with pdfplumber.open(path) as document:
        for index, page in enumerate(document.pages, start=1):
            text = (page.extract_text() or "").strip()
            if len(text) < SCANNED_TEXT_THRESHOLD:
                scanned.append(index)
            else:
                chunks.append(Chunk(text=text, file=path.name, page=index))
    return chunks, _render_pages(path, scanned)


def _render_pages(path: Path, pages: list[int]) -> list[PageImage]:
    if not pages:
        return []
    document = pdfium.PdfDocument(str(path))
    try:
        return [
            PageImage(
                data=_encode(document[number - 1].render(scale=RENDER_SCALE).to_pil()),
                media_type="image/png",
                file=path.name,
                page=number,
            )
            for number in pages
        ]
    finally:
        document.close()


def _load_image(path: Path) -> tuple[list[Chunk], list[PageImage]]:
    """An uploaded photo or scan. No text layer exists, so it is vision-only."""
    with Image.open(path) as image:
        payload = _encode(image.convert("RGB"))
    return [], [PageImage(data=payload, media_type="image/png", file=path.name, page=1)]


def _encode(image: Image.Image) -> bytes:
    """PNG bytes, downscaled if oversized.

    A full-resolution scan carries no more readable detail than a 2000px one but
    costs several times as much to send.
    """
    if max(image.size) > MAX_IMAGE_EDGE:
        ratio = MAX_IMAGE_EDGE / max(image.size)
        image = image.resize((int(image.width * ratio), int(image.height * ratio)))
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _load_docx(path: Path) -> tuple[list[Chunk], list[PageImage]]:
    """Paragraphs and tables. Tables carry the terms, so they cannot be skipped."""
    document = Docx(path)
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        lines.extend(
            " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
        )
    text = "\n".join(lines)
    return ([Chunk(text=text, file=path.name)] if text else [], [])


def _load_excel(path: Path) -> tuple[list[Chunk], list[PageImage]]:
    """One chunk per sheet, each cell tagged with its reference.

    The coordinates are the point: a citation reading "B7 of the Quotation sheet" is
    checkable, where "somewhere in the spreadsheet" is not.
    """
    workbook = load_workbook(path, data_only=True)
    chunks: list[Chunk] = []
    for sheet in workbook.worksheets:
        lines = [
            " | ".join(
                f"{cell.coordinate}={cell.value}"
                for cell in row
                if cell.value not in (None, "")
            )
            for row in sheet.iter_rows()
        ]
        body = "\n".join(line for line in lines if line)
        if body:
            chunks.append(
                Chunk(text=body, file=path.name, sheet=sheet.title)
            )
    workbook.close()
    return chunks, []


def _load_delimited(path: Path) -> tuple[list[Chunk], list[PageImage]]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with path.open(newline="", encoding="utf-8-sig") as handle:
        rows = list(csv.reader(handle, delimiter=delimiter))
    lines = [
        f"row {number}: " + " | ".join(cell.strip() for cell in row)
        for number, row in enumerate(rows, start=1)
        if any(cell.strip() for cell in row)
    ]
    text = "\n".join(lines)
    return ([Chunk(text=text, file=path.name)] if text else [], [])


def _load_text(path: Path) -> tuple[list[Chunk], list[PageImage]]:
    text = path.read_text(encoding="utf-8-sig").strip()
    return ([Chunk(text=text, file=path.name)] if text else [], [])
